"""文档解析服务

支持解析多种文档格式，提取文本内容供 LLM 使用。
支持的格式：PDF, DOCX, TXT, MD, JSON, CSV, 代码文件等。
"""

import io
from typing import Optional, Tuple, List, Dict, Any, Union
from loguru import logger
import httpx


# ========== 文档解析器 ==========

# 内容块类型：文本或图片
ContentBlock = Dict[str, Any]  # {"type": "text", "content": str, "page": int} 或 {"type": "image", "data": bytes, "page": int}


def parse_pdf_structured(file_bytes: bytes, extract_images: bool = True) -> List[ContentBlock]:
    """解析 PDF 文档，按页顺序返回文本和图片（保留位置关系）
    
    Args:
        file_bytes: PDF 文件字节
        extract_images: 是否提取图片
        
    Returns:
        内容块列表，按页顺序排列，每页先文本后图片
    """
    content_blocks: List[ContentBlock] = []
    
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        for i, page in enumerate(doc, 1):
            # 先添加该页的文本
            page_text = page.get_text()
            if page_text and page_text.strip():
                content_blocks.append({
                    "type": "text",
                    "content": f"--- 第 {i} 页 ---\n{page_text}",
                    "page": i
                })
            
            # 再添加该页的图片（图片紧跟在该页文本后面）
            if extract_images:
                image_list = page.get_images(full=True)
                for img_index, img_info in enumerate(image_list):
                    try:
                        xref = img_info[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # 过滤太小的图片（可能是图标/装饰）
                        if len(image_bytes) > 5000:  # > 5KB
                            content_blocks.append({
                                "type": "image",
                                "data": image_bytes,
                                "page": i,
                                "index": img_index
                            })
                            logger.debug(f"[DocumentParser] 提取图片: 页{i}, 索引{img_index}, {len(image_bytes)} 字节")
                    except Exception as e:
                        logger.warning(f"[DocumentParser] 图片提取失败: 页{i}, 索引{img_index}, {e}")
        
        page_count = len(doc)
        image_count = len([b for b in content_blocks if b["type"] == "image"])
        text_count = len([b for b in content_blocks if b["type"] == "text"])
        doc.close()
        
        logger.info(f"[DocumentParser] PDF 结构化解析完成 (PyMuPDF): {page_count} 页, {text_count} 个文本块, {image_count} 张图片")
        return content_blocks
        
    except ImportError:
        logger.warning("[DocumentParser] PyMuPDF 未安装，回退到 pypdf（不支持图片提取）")
    except Exception as e:
        logger.warning(f"[DocumentParser] PyMuPDF 解析失败: {e}，回退到 pypdf")
    
    # 回退方案：使用 pypdf（不支持图片提取）
    try:
        from pypdf import PdfReader
        
        reader = PdfReader(io.BytesIO(file_bytes))
        
        for i, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                content_blocks.append({
                    "type": "text",
                    "content": f"--- 第 {i} 页 ---\n{page_text}",
                    "page": i
                })
        
        logger.info(f"[DocumentParser] PDF 结构化解析完成 (pypdf): {len(reader.pages)} 页, 无图片")
        return content_blocks
        
    except ImportError:
        logger.error("[DocumentParser] PDF 解析库未安装，请运行: pip install pymupdf pypdf")
        return [{"type": "text", "content": "[错误: PDF 解析库未安装]", "page": 0}]
    except Exception as e:
        logger.error(f"[DocumentParser] PDF 解析失败: {e}")
        return [{"type": "text", "content": f"[PDF 解析失败: {e}]", "page": 0}]


def parse_pdf(file_bytes: bytes, extract_images: bool = True) -> Tuple[str, List[bytes]]:
    """解析 PDF 文档，提取文本和图片（兼容旧接口）
    
    Args:
        file_bytes: PDF 文件字节
        extract_images: 是否提取图片
        
    Returns:
        (text, images): 文本内容和提取的图片列表（bytes）
    """
    blocks = parse_pdf_structured(file_bytes, extract_images)
    
    text_parts = [b["content"] for b in blocks if b["type"] == "text"]
    images = [b["data"] for b in blocks if b["type"] == "image"]
    
    return "\n\n".join(text_parts), images


def parse_docx(file_bytes: bytes) -> str:
    """解析 Word 文档 (.docx)
    
    Args:
        file_bytes: DOCX 文件字节
        
    Returns:
        文本内容
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            if table_rows:
                text_parts.append("\n[表格]\n" + "\n".join(table_rows))
        
        text = "\n\n".join(text_parts)
        logger.info(f"[DocumentParser] DOCX 解析完成: {len(text)} 字符")
        return text
        
    except ImportError:
        logger.error("[DocumentParser] python-docx 未安装，请运行: pip install python-docx")
        return "[错误: DOCX 解析库未安装]"
    except Exception as e:
        logger.error(f"[DocumentParser] DOCX 解析失败: {e}")
        return f"[DOCX 解析失败: {e}]"


def parse_excel(file_bytes: bytes, filename: str) -> str:
    """解析 Excel 文件 (.xlsx, .xls)
    
    Args:
        file_bytes: Excel 文件字节
        filename: 文件名（用于判断格式）
        
    Returns:
        文本内容（表格形式）
    """
    try:
        import pandas as pd
        
        if filename.lower().endswith('.xlsx'):
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes), engine='openpyxl')
        else:
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes), engine='xlrd')
        
        text_parts = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            if not df.empty:
                # 转换为 Markdown 表格格式
                table_md = df.to_markdown(index=False)
                text_parts.append(f"### Sheet: {sheet_name}\n\n{table_md}")
        
        text = "\n\n".join(text_parts)
        logger.info(f"[DocumentParser] Excel 解析完成: {len(excel_file.sheet_names)} 个工作表, {len(text)} 字符")
        return text
        
    except ImportError as e:
        logger.error(f"[DocumentParser] Excel 解析库未安装: {e}")
        return "[错误: Excel 解析库未安装，请安装 pandas, openpyxl]"
    except Exception as e:
        logger.error(f"[DocumentParser] Excel 解析失败: {e}")
        return f"[Excel 解析失败: {e}]"


def parse_text_file(file_bytes: bytes, encoding: str = 'utf-8') -> str:
    """解析纯文本文件 (.txt, .md, .json, .log, 代码文件等)
    
    Args:
        file_bytes: 文件字节
        encoding: 编码格式
        
    Returns:
        文本内容
    """
    try:
        # 尝试多种编码
        for enc in [encoding, 'utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                text = file_bytes.decode(enc)
                logger.info(f"[DocumentParser] 文本解析完成 (编码: {enc}): {len(text)} 字符")
                return text
            except UnicodeDecodeError:
                continue
        
        # 最后尝试忽略错误
        text = file_bytes.decode('utf-8', errors='ignore')
        logger.warning(f"[DocumentParser] 文本解析使用 ignore 模式: {len(text)} 字符")
        return text
        
    except Exception as e:
        logger.error(f"[DocumentParser] 文本解析失败: {e}")
        return f"[文本解析失败: {e}]"


def parse_csv(file_bytes: bytes) -> str:
    """解析 CSV 文件
    
    Args:
        file_bytes: CSV 文件字节
        
    Returns:
        文本内容（表格形式）
    """
    try:
        import pandas as pd
        
        # 尝试不同编码
        for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                df = pd.read_csv(io.BytesIO(file_bytes), encoding=enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            df = pd.read_csv(io.BytesIO(file_bytes), encoding='utf-8', errors='ignore')
        
        # 限制行数，避免内容过长
        if len(df) > 100:
            text = f"[CSV 共 {len(df)} 行，显示前 100 行]\n\n"
            text += df.head(100).to_markdown(index=False)
        else:
            text = df.to_markdown(index=False)
        
        logger.info(f"[DocumentParser] CSV 解析完成: {len(df)} 行, {len(text)} 字符")
        return text
        
    except ImportError:
        logger.error("[DocumentParser] pandas 未安装")
        return "[错误: CSV 解析库未安装]"
    except Exception as e:
        logger.error(f"[DocumentParser] CSV 解析失败: {e}")
        return f"[CSV 解析失败: {e}]"


# ========== 主解析函数 ==========

def parse_document(file_bytes: bytes, content_type: str, filename: str) -> Tuple[str, List[str]]:
    """解析文档，提取文本内容
    
    Args:
        file_bytes: 文件字节
        content_type: MIME 类型
        filename: 文件名
        
    Returns:
        (text, images): 文本内容和提取的图片列表
    """
    filename_lower = filename.lower()
    
    # PDF
    if content_type == 'application/pdf' or filename_lower.endswith('.pdf'):
        return parse_pdf(file_bytes)
    
    # Word
    if content_type in [
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ] or filename_lower.endswith(('.docx', '.doc')):
        return parse_docx(file_bytes), []
    
    # Excel
    if content_type in [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel'
    ] or filename_lower.endswith(('.xlsx', '.xls')):
        return parse_excel(file_bytes, filename), []
    
    # CSV
    if content_type == 'text/csv' or filename_lower.endswith('.csv'):
        return parse_csv(file_bytes), []
    
    # 纯文本类（txt, md, json, log, 代码文件等）
    text_extensions = {
        '.txt', '.md', '.markdown', '.json', '.log',
        '.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs', '.rb', '.php',
        '.html', '.css', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg',
        '.sh', '.bash', '.zsh', '.sql'
    }
    
    if content_type.startswith('text/') or any(filename_lower.endswith(ext) for ext in text_extensions):
        return parse_text_file(file_bytes), []
    
    # JSON (可能是 application/json)
    if content_type == 'application/json' or filename_lower.endswith('.json'):
        return parse_text_file(file_bytes), []
    
    # 未知类型，尝试作为文本解析
    logger.warning(f"[DocumentParser] 未知文档类型: {content_type}, {filename}，尝试作为文本解析")
    return parse_text_file(file_bytes), []


async def parse_document_from_url(url: str, content_type: str, filename: str) -> Tuple[str, List[bytes]]:
    """从 URL 下载并解析文档
    
    Args:
        url: 文档 URL
        content_type: MIME 类型
        filename: 文件名
        
    Returns:
        (text, images): 文本内容和提取的图片列表
    """
    try:
        # 下载文件
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            file_bytes = response.content
        
        logger.info(f"[DocumentParser] 下载文档完成: {filename}, {len(file_bytes)} 字节")
        
        # 解析文档
        return parse_document(file_bytes, content_type, filename)
        
    except httpx.TimeoutException:
        logger.error(f"[DocumentParser] 下载超时: {url}")
        return f"[文档下载超时: {filename}]", []
    except httpx.HTTPError as e:
        logger.error(f"[DocumentParser] 下载失败: {url}, {e}")
        return f"[文档下载失败: {filename}]", []
    except Exception as e:
        logger.error(f"[DocumentParser] 解析失败: {url}, {e}")
        return f"[文档解析失败: {filename}, {e}]", []


async def parse_document_structured_from_url(url: str, content_type: str, filename: str) -> List[ContentBlock]:
    """从 URL 下载并结构化解析文档（保留图片位置关系）
    
    Args:
        url: 文档 URL
        content_type: MIME 类型
        filename: 文件名
        
    Returns:
        内容块列表，按页顺序排列
    """
    try:
        # 下载文件
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            file_bytes = response.content
        
        logger.info(f"[DocumentParser] 下载文档完成: {filename}, {len(file_bytes)} 字节")
        
        filename_lower = filename.lower()
        
        # PDF - 使用结构化解析
        if content_type == 'application/pdf' or filename_lower.endswith('.pdf'):
            return parse_pdf_structured(file_bytes)
        
        # 其他格式 - 返回纯文本块
        text, _ = parse_document(file_bytes, content_type, filename)
        return [{"type": "text", "content": text, "page": 0}]
        
    except httpx.TimeoutException:
        logger.error(f"[DocumentParser] 下载超时: {url}")
        return [{"type": "text", "content": f"[文档下载超时: {filename}]", "page": 0}]
    except httpx.HTTPError as e:
        logger.error(f"[DocumentParser] 下载失败: {url}, {e}")
        return [{"type": "text", "content": f"[文档下载失败: {filename}]", "page": 0}]
    except Exception as e:
        logger.error(f"[DocumentParser] 解析失败: {url}, {e}")
        return [{"type": "text", "content": f"[文档解析失败: {filename}, {e}]", "page": 0}]


# ========== 工具函数 ==========

def truncate_text(text: str, max_chars: int = 50000) -> str:
    """截断过长的文本
    
    Args:
        text: 原始文本
        max_chars: 最大字符数
        
    Returns:
        截断后的文本
    """
    if len(text) <= max_chars:
        return text
    
    # 截断并添加提示
    truncated = text[:max_chars]
    return f"{truncated}\n\n... [文档内容过长，已截断，共 {len(text)} 字符，显示前 {max_chars} 字符]"
